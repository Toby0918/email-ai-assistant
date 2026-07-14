"""Bounded local text extraction for already validated private attachments."""

from __future__ import annotations

from pathlib import Path

from .attachment_security import DOCX, JPEG, PDF, PNG, TIFF, XLSX


MAX_TEXT_CHARACTERS = 2_000_000
MAX_SHEET_CELLS = 100_000


def parse_private_attachment(path: Path, mime_type: str) -> bytes:
    if mime_type == PDF:
        text = _pdf(path)
    elif mime_type == DOCX:
        text = _docx(path)
    elif mime_type == XLSX:
        text = _xlsx(path)
    elif mime_type in {PNG, JPEG, TIFF}:
        text = _image(path)
    else:
        raise ValueError("unsupported_attachment")
    if len(text) > MAX_TEXT_CHARACTERS:
        raise ValueError("attachment_text_limit")
    return text.encode("utf-8", errors="strict")


def _pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path), strict=True)
    if reader.is_encrypted or len(reader.pages) > 500:
        raise ValueError("pdf_unsupported")
    pieces: list[str] = []
    total = 0
    for page in reader.pages:
        value = page.extract_text() or ""
        total += len(value)
        if total > MAX_TEXT_CHARACTERS:
            raise ValueError("attachment_text_limit")
        pieces.append(value)
    return "\n".join(pieces)


def _docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    pieces = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            pieces.append("\t".join(cell.text for cell in row.cells))
    result = "\n".join(pieces)
    if len(result) > MAX_TEXT_CHARACTERS:
        raise ValueError("attachment_text_limit")
    return result


def _xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True, keep_links=False)
    pieces: list[str] = []
    cells = 0
    try:
        for worksheet in workbook.worksheets:
            for row in worksheet.iter_rows(values_only=True):
                cells += len(row)
                if cells > MAX_SHEET_CELLS:
                    raise ValueError("spreadsheet_cell_limit")
                pieces.append("\t".join("" if value is None else str(value) for value in row))
                if sum(len(item) for item in pieces) > MAX_TEXT_CHARACTERS:
                    raise ValueError("attachment_text_limit")
    finally:
        workbook.close()
    return "\n".join(pieces)


def _image(path: Path) -> str:
    try:
        import pytesseract
        from PIL import Image

        with Image.open(path) as image:
            image.verify()
        with Image.open(path) as image:
            if image.width * image.height > 40_000_000:
                raise ValueError("image_pixel_limit")
            return pytesseract.image_to_string(image)
    except (ImportError, OSError):
        return ""


__all__ = ["parse_private_attachment"]
