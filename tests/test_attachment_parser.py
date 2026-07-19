"""Tests for bounded, de-identified attachment parsing."""

from __future__ import annotations

import json
import inspect
import sqlite3
import struct
import time
import unittest
from datetime import UTC, datetime, timedelta
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import MagicMock, patch
from zipfile import ZIP_DEFLATED, ZIP_STORED, ZipFile

import pypdf.filters as pypdf_filters
from docx import Document
from openpyxl import Workbook
from PIL import Image

from backend.email_agent import attachment_parser
from backend.email_agent import attachment_safety
from backend.email_agent.analyzer import build_analysis_prompt
from backend.email_agent.attachment_fact_safety import sanitize_constructed_fact
from backend.email_agent.attachment_facts import extract_attachment_facts
from backend.email_agent.attachment_parser import (
    bind_prepared_media_evidence,
    parse_attachment_bundles_compat,
    parse_attachments,
)
from backend.email_agent.attachment_model_context import (
    AttachmentAnalysisBundle,
    attachment_model_candidate,
)
from backend.email_agent.attachment_media_context import (
    UNTRUSTED_MEDIA_EVIDENCE,
    provider_attachment_candidate,
)
from backend.email_agent.attachment_storage import StoredAttachment
from backend.email_agent.attachment_text import sanitize_text
from backend.email_agent.database import initialize_schema, save_analysis
from backend.email_agent.multimodal_media import PreparedMediaAsset


EXPECTED_KEYS = {"filename", "type", "status", "summary", "key_facts", "limitations"}


class _ObservedParagraph:
    def __init__(self, text: str) -> None:
        self._text = text
        self.read_count = 0

    @property
    def text(self) -> str:
        self.read_count += 1
        return self._text


class _ObservedCell:
    def __init__(self, text: str) -> None:
        self.text = text
        self.read_count = 0

    def __str__(self) -> str:
        self.read_count += 1
        return self.text


class _ObservedDocxCell:
    def __init__(self, text: str) -> None:
        self._text = text
        self.read_count = 0

    @property
    def text(self) -> str:
        self.read_count += 1
        return self._text


class AttachmentParserTests(unittest.TestCase):
    def test_successful_media_without_text_gets_one_fixed_untrusted_media_candidate(self) -> None:
        metadata = AttachmentAnalysisBundle(
            {
                "filename": "visible.png",
                "type": "image",
                "status": "metadata_only",
                "summary": "Image metadata only.",
                "key_facts": [],
                "limitations": ["OCR returned no readable text."],
            },
            None,
        )
        asset = PreparedMediaAsset(
            source_id="attachment:0",
            provider_filename="image_0.png",
            mime_type="image/png",
            kind="image",
            detail="high",
            buffer=bytearray(b"synthetic"),
        )

        bound = bind_prepared_media_evidence((metadata,), (asset,))

        self.assertEqual(bound[0].display_insight, metadata.display_insight)
        self.assertEqual(bound[0].model_candidate.source_id, "attachment:0")
        self.assertEqual(
            bound[0].model_candidate.text,
            "UNTRUSTED_MEDIA: sanitized current-message media has no locally extracted text.",
        )
        self.assertTrue(bound[0].model_candidate.visual_only)
        self.assertIsNone(
            provider_attachment_candidate(bound[0].model_candidate, "deepseek")
        )

    def test_marker_text_without_visual_provenance_remains_a_text_candidate(self) -> None:
        candidate = attachment_model_candidate(
            "attachment:0", UNTRUSTED_MEDIA_EVIDENCE,
        )

        self.assertFalse(candidate.visual_only)
        self.assertIs(provider_attachment_candidate(candidate, "deepseek"), candidate)
        self.assertIs(provider_attachment_candidate(candidate, "openai"), candidate)

    def test_failed_media_gets_no_candidate_and_text_candidate_is_preserved(self) -> None:
        metadata = AttachmentAnalysisBundle({"status": "metadata_only"}, None)
        existing = attachment_model_candidate("attachment:1", "bounded extracted text")
        parsed = AttachmentAnalysisBundle({"status": "parsed"}, existing)

        bound = bind_prepared_media_evidence((metadata, parsed), ())

        self.assertIsNone(bound[0].model_candidate)
        self.assertIs(bound[1].model_candidate, existing)

    def test_parsed_bundle_has_private_candidate_and_metadata_only_does_not(self) -> None:
        raw = "PO 1013970520 qty 24 due 2026-07-20 https://private.example.test/a"
        with TemporaryDirectory() as directory:
            parsed = self._write(directory, "parsed.pdf", "pdf", b"synthetic")
            metadata_only = self._write(directory, "wrong.txt", "pdf", b"synthetic")
            page = MagicMock()
            page.extract_text.return_value = raw
            reader = MagicMock()
            reader.pages = [page]

            with patch.object(attachment_parser, "PdfReader", return_value=reader):
                bundles = parse_attachment_bundles_compat([parsed, metadata_only])

        self.assertEqual(bundles[0].display_insight["status"], "parsed")
        self.assertIsNotNone(bundles[0].model_candidate)
        self.assertEqual(bundles[0].model_candidate.source_id, "attachment:0")
        self.assertIn("PO 1013970520", bundles[0].model_candidate.text)
        self.assertIsNone(bundles[1].model_candidate)
        self.assertFalse(hasattr(bundles[0], "__dict__"))
        self.assertNotIn(raw, repr(bundles[0]))

    def test_public_attachment_output_recursively_excludes_model_projection(self) -> None:
        raw = "PO 1013970520 qty 24 due 2026-07-20 PRIVATE-MODEL-TEXT"
        with TemporaryDirectory() as directory:
            parsed = self._write(directory, "parsed.pdf", "pdf", b"synthetic")
            page = MagicMock()
            page.extract_text.return_value = raw
            reader = MagicMock()
            reader.pages = [page]

            with patch.object(attachment_parser, "PdfReader", return_value=reader):
                public_output = parse_attachments([parsed])

        serialized = json.dumps(public_output, ensure_ascii=False)
        self.assertEqual(set(public_output[0]), EXPECTED_KEYS)
        for forbidden in (
            "model_candidate",
            "model_text",
            "source_id",
            "attachment:0",
            "PRIVATE-MODEL-TEXT",
        ):
            with self.subTest(forbidden=forbidden):
                self.assertNotIn(forbidden, serialized)

    def test_generic_sanitizer_redacts_every_long_digit_sequence(self) -> None:
        cases = (
            "RFQ7654321",
            "PO123456789012",
            "accountABC1234567890",
            "acctABC1234-5678-9012-3456XYZ",
            "due 2026-07-18",
            "RFQ1234 567",
            "RFQ1234\t567",
            "acctABC1234 5678 9012 3456XYZ",
            r"path1234\567",
            "group1234_567",
            "group1234/567",
            "group1234:567",
            "group1234+567",
            "group1234.567",
        )

        for value in cases:
            with self.subTest(value=value):
                sanitized = sanitize_text(value)
                self.assertLess(
                    sum(character.isdigit() for character in sanitized),
                    7,
                    sanitized,
                )

    def test_identifier_gate_rejects_prefixed_sensitive_shapes_and_deduplicates(self) -> None:
        facts = extract_attachment_facts("\n".join([
            "RFQ: 202/555/0199",
            "Reference: RFQ-202-555-0199",
            "Invoice: INV-1234-5678-9012",
            "Reference: RFQ-/home/private/1234",
            "Reference: RFQ-www.private.example/1234",
            "Reference: RFQ-private.example/1234",
            "Reference: RFQ-RFQ-202-555-0199",
            "Reference: RFQ-X202-555-0199",
            "RFQ: 13800138000",
            "RFQ: 2025550199",
            "Invoice: INV-7000001",
            "Reference: INV-7000001",
        ]))

        self.assertEqual(facts, ["Reference: Invoice INV-7000001"])

    def test_negated_or_resolved_attachment_statements_do_not_become_facts(self) -> None:
        facts = extract_attachment_facts("\n".join([
            "This item is not due 2026-07-18.",
            "Do not respond within 3 days.",
            "Please do not confirm quantity.",
            "Action required: never provide quotation.",
            "The part is not currently damaged.",
            "There is no evidence of burrs.",
            "The dimension is not currently out of tolerance.",
            "The prior surface damage is now resolved.",
            "No scratches were found.",
            "Please don't confirm quantity.",
            "The part isn't damaged.",
            "The item is free of burrs.",
            "Due date: 2026-07-20.",
            "Please provide the quotation.",
            "Quality issue: scratched surface.",
        ]))

        self.assertEqual(
            facts,
            [
                "Deadline: due 2026-07-20",
                "Requested action: provide quotation",
                "Quality issue: surface_damage",
            ],
        )

    def test_modal_contractions_do_not_become_attachment_facts(self) -> None:
        facts = extract_attachment_facts("\n".join([
            "The item won't be due 2026-07-21.",
            "Please note that we wouldn't provide quotation.",
            "Please note that we shouldn't confirm quantity.",
            "The part couldn't be damaged.",
            "Please note that we mustn't review specification.",
            "The item won’t be due 2026-07-22.",
            "The item cannot be due 2026-07-23.",
            "The item shan't be due 2026-07-24.",
            "The item oughtn't be due 2026-07-25.",
        ]))

        self.assertEqual(facts, [])

    def test_identifier_extraction_requires_a_complete_consistent_field(self) -> None:
        rejected = (
            "PO: 021-000-021",
            "PO: 123-45-6789",
            "Invoice: 12-3456789",
            "RFQ: 202555019 9",
            "RFQ: 138001380 00",
            "RFQ: 4111 1111 1111 1111",
            "RFQ: RFQRFQ7654321",
            "RFQ: PO-7654322",
            "Reference: RFQ-RFQ7654323",
            "RFQ: 7654321 extra",
            r"RFQ: X1234\private",
            "Invoice: RFQ-7654321",
            "RFQ | 202555019 | 9",
            "RFQ | 138001380 | 00",
            "RFQ | 4111 | 1111 | 1111 | 1111",
        )
        for value in rejected:
            with self.subTest(value=value):
                self.assertEqual(extract_attachment_facts(value), [])

        accepted = {
            "RFQ: 7654321.": ["Reference: RFQ 7654321"],
            "RFQ: ABC1234": ["Reference: RFQ ABC1234"],
            "Reference: RFQ-SAFE_42": ["Reference: RFQ-SAFE_42"],
        }
        for value, expected in accepted.items():
            with self.subTest(value=value):
                self.assertEqual(extract_attachment_facts(value), expected)

    def test_multicell_identifier_continuations_never_cross_boundaries(self) -> None:
        document = Document()
        table = document.add_table(rows=3, cols=5)
        docx_rows = (
            ("RFQ", "202555019", "9"),
            ("RFQ", "138001380", "00"),
            ("RFQ", "4111", "1111", "1111", "1111"),
        )
        for row_index, values in enumerate(docx_rows):
            for cell_index, value in enumerate(values):
                table.cell(row_index, cell_index).text = value
        docx_content = BytesIO()
        document.save(docx_content)

        workbook = Workbook()
        sheet = workbook.active
        for values in docx_rows:
            sheet.append(values)
        xlsx_content = BytesIO()
        workbook.save(xlsx_content)

        with TemporaryDirectory() as directory:
            insights = parse_attachments([
                self._write(directory, "continuation.docx", "docx", docx_content.getvalue()),
                self._write(directory, "continuation.xlsx", "xlsx", xlsx_content.getvalue()),
            ])

        for insight in insights:
            with self.subTest(filename=insight["filename"]):
                self.assertEqual(insight["key_facts"], [])

        prompt = build_analysis_prompt(
            subject="Synthetic table continuation",
            sender="sender@example.test",
            clean_body="Please review the synthetic table.",
            attachment_insights=insights,
        )
        connection = sqlite3.connect(":memory:")
        initialize_schema(connection)
        save_analysis(
            connection,
            subject="Synthetic table continuation",
            sender="sender@example.test",
            analysis={"summary": "Safe synthetic result.", "attachment_insights": insights},
        )
        stored_json = connection.execute(
            "SELECT analysis_json FROM email_analysis"
        ).fetchone()[0]
        serialized_result = json.dumps(insights, ensure_ascii=False)
        for secret in ("202555019", "138001380", "Reference: RFQ 4111"):
            for boundary, value in (
                ("result", serialized_result),
                ("prompt", prompt),
                ("storage", stored_json),
            ):
                with self.subTest(boundary=boundary, secret=secret):
                    self.assertNotIn(secret, value)

    def test_identifier_deduplication_uses_complete_canonical_identity(self) -> None:
        cases = {
            "RFQ: ABC1234\nPO: 1234": [
                "Reference: RFQ ABC1234",
                "Reference: PO 1234",
            ],
            "RFQ: XABC1234\nPO: ABC1234": [
                "Reference: RFQ XABC1234",
                "Reference: PO ABC1234",
            ],
            "Invoice: INV-7000001\nReference: INV-7000001": [
                "Reference: Invoice INV-7000001",
            ],
        }
        for value, expected in cases.items():
            with self.subTest(value=value):
                self.assertEqual(extract_attachment_facts(value), expected)

    def test_reversal_context_is_bounded_to_the_candidate_clause(self) -> None:
        rejected = (
            "It won't be due 2026-07-18.",
            "Action required: you shouldn't confirm quantity.",
            "It hasn’t been due 2026-07-19.",
            "It hadn't been due 2026-07-20.",
            "The item is free from burrs.",
            "Burrs are absent.",
            "The scratch was repaired.",
            "Leakage stopped.",
            "Burrs were eliminated.",
            "The damage was remediated.",
            "Deadline: 2026-07-18 was withdrawn.",
            "Deadline: 2026-07-18 was waived.",
            "Deadline: 2026-07-18 was cancelled.",
            "Deadline: 2026-07-18 was revoked.",
            "Action required: cancel the request to confirm quantity.",
            "Please disregard the request to provide quotation.",
            "Please skip the request to review specification.",
        )
        for value in rejected:
            with self.subTest(value=value):
                self.assertEqual(extract_attachment_facts(value), [])

        accepted = {
            "Due date: 2026-07-18, but do not miss it.": [
                "Deadline: due 2026-07-18",
            ],
            "Please confirm quantity, not price.": [
                "Requested action: confirm quantity",
            ],
            "damaged but not leaking": [
                "Quality issue: physical_damage",
            ],
            "scratched, but no burrs": [
                "Quality issue: surface_damage",
            ],
        }
        for value, expected in accepted.items():
            with self.subTest(value=value):
                self.assertEqual(extract_attachment_facts(value), expected)

    def test_later_positive_action_survives_an_earlier_reversed_clause(self) -> None:
        cases = {
            "Please do not confirm price, but confirm quantity.": [
                "Requested action: confirm quantity",
            ],
            "Please disregard the request to provide quotation, but confirm quantity.": [
                "Requested action: confirm quantity",
            ],
            "Please don't provide quotation, however review specification.": [
                "Requested action: review specification",
            ],
            "Please shouldn’t confirm price, provide quotation.": [
                "Requested action: provide quotation",
            ],
            "Action required: skip the request to review specification; however confirm quantity.": [
                "Requested action: confirm quantity",
            ],
            "Please cancel quotation and confirm quantity.": [
                "Requested action: confirm quantity",
            ],
            "Please decline quotation, then confirm quantity.": [
                "Requested action: confirm quantity",
            ],
            "Please archive invoice and confirm quantity.": [
                "Requested action: confirm quantity",
            ],
            "Please resolve damaged part but confirm quantity.": [
                "Requested action: confirm quantity",
            ],
            "Please confirm not price but quantity.": [
                "Requested action: confirm quantity",
            ],
            "Please confirm no quotation but quantity.": [
                "Requested action: confirm quantity",
            ],
        }
        for value, expected in cases.items():
            with self.subTest(value=value):
                self.assertEqual(extract_attachment_facts(value), expected)

    def test_later_positive_action_reaches_prompt_and_storage_without_reversed_text(self) -> None:
        payload = (
            "Please disregard the request to provide quotation, but confirm quantity."
        )
        with TemporaryDirectory() as directory:
            item = self._write(directory, "action.pdf", "pdf", b"synthetic")
            page = MagicMock()
            page.extract_text.return_value = payload
            reader = MagicMock()
            reader.pages = [page]
            with patch.object(attachment_parser, "PdfReader", return_value=reader):
                insights = parse_attachments([item])

        self.assertEqual(
            insights[0]["key_facts"],
            ["Requested action: confirm quantity"],
        )
        prompt = build_analysis_prompt(
            subject="Synthetic reversed action",
            sender="sender@example.test",
            clean_body="Please review the attachment.",
            attachment_insights=insights,
        )
        connection = sqlite3.connect(":memory:")
        initialize_schema(connection)
        save_analysis(
            connection,
            subject="Synthetic reversed action",
            sender="sender@example.test",
            analysis={"summary": "Safe synthetic result.", "attachment_insights": insights},
        )
        stored_json = connection.execute(
            "SELECT analysis_json FROM email_analysis"
        ).fetchone()[0]
        serialized = json.dumps(insights, ensure_ascii=False)
        for boundary, value in (
            ("result", serialized),
            ("prompt", prompt),
            ("storage", stored_json),
        ):
            with self.subTest(boundary=boundary):
                self.assertIn("Requested action: confirm quantity", value)
                self.assertNotIn("Requested action: provide quotation", value)
                self.assertNotIn("disregard the request", value)

    def test_completed_action_does_not_leak_its_verb_into_an_unsupported_clause(self) -> None:
        cases = {
            "Please confirm quantity, but cancel quotation.": [
                "Requested action: confirm quantity",
            ],
            "Please confirm quantity and archive invoice.": [
                "Requested action: confirm quantity",
            ],
        }
        for value, expected in cases.items():
            with self.subTest(value=value):
                self.assertEqual(extract_attachment_facts(value), expected)

    def test_adjacent_quality_absence_does_not_become_a_positive_fact(self) -> None:
        rejected = (
            "damage-free",
            "damage free",
            "scratch-free",
            "scratch free",
            "leak-free",
            "leak free",
            "burr-free",
            "burr free",
            "zero scratches",
            "0 scratches",
            "0 burrs",
            "zero leakage",
            "nil scratches",
            "non-damaged",
            "non-scratched",
        )
        for value in rejected:
            with self.subTest(value=value):
                self.assertEqual(extract_attachment_facts(value), [])

        self.assertEqual(
            extract_attachment_facts("noncritical damage"),
            ["Quality issue: physical_damage"],
        )

    def test_deadline_post_context_distinguishes_absence_from_negated_retirement(self) -> None:
        rejected = (
            "Due date: 2026-07-18 is not required.",
            "Deadline: 2026-07-19 does not apply.",
            "Deadline: 2026-07-20 is no longer applicable.",
            "Deadline: 2026-07-21 is optional.",
        )
        for value in rejected:
            with self.subTest(value=value):
                self.assertEqual(extract_attachment_facts(value), [])

        self.assertEqual(
            extract_attachment_facts(
                "Deadline: 2026-07-22 remains active and is not waived."
            ),
            ["Deadline: due 2026-07-22"],
        )

    def test_quality_resolution_context_preserves_explicitly_unresolved_issues(self) -> None:
        accepted = {
            "The scratch was not repaired.": ["Quality issue: surface_damage"],
            "The damage is not resolved.": ["Quality issue: physical_damage"],
            "Burrs have not been removed.": ["Quality issue: burrs"],
            "Leakage has not stopped.": ["Quality issue: leakage"],
        }
        for value, expected in accepted.items():
            with self.subTest(value=value):
                self.assertEqual(extract_attachment_facts(value), expected)

        rejected = (
            "The scratch was repaired.",
            "The damage is resolved.",
            "Burrs have been removed.",
            "Leakage has stopped.",
        )
        for value in rejected:
            with self.subTest(value=value):
                self.assertEqual(extract_attachment_facts(value), [])

    def test_cross_format_sensitive_identifiers_and_negations_never_cross_boundaries(self) -> None:
        payload = "\n".join([
            "RFQ: 202/555/0199",
            "Reference: RFQ-202-555-0199",
            "Invoice: INV-1234-5678-9012",
            "Reference: RFQ-/home/private/1234",
            "Reference: RFQ-www.private.example/1234",
            "Reference: RFQ-private.example/1234",
            "Reference: RFQ-RFQ-202-555-0199",
            "RFQ: 13800138000",
            "This item is not due 2026-07-18.",
            "Do not respond within 3 days.",
            "Please do not confirm quantity.",
            "Action required: never provide quotation.",
            "The part is not currently damaged.",
            "There is no evidence of burrs.",
            "The dimension is not currently out of tolerance.",
            "No scratches were found.",
            "Please don't confirm quantity.",
            "The part isn't damaged.",
            "PO: 021-000-021",
            "PO: 123-45-6789",
            "Invoice: 12-3456789",
            "RFQ: 202555019 9",
            "RFQ: 138001380 00",
            "RFQ: 4111 1111 1111 1111",
            "RFQ: RFQRFQ7654321",
            "RFQ: PO-7654322",
            "Reference: RFQ-RFQ7654323",
            "It won't be due 2026-07-19.",
            "Action required: you shouldn't confirm quantity.",
            "Burrs are absent.",
            "The scratch was repaired.",
            "Deadline: 2026-07-20 was revoked.",
            "Please disregard the request to provide quotation.",
            "RFQ: 7654321",
        ])
        with TemporaryDirectory() as directory:
            items = [
                self._write(directory, "security.pdf", "pdf", b"synthetic"),
                self._write(directory, "security.docx", "docx", self._security_docx_bytes(payload)),
                self._write(directory, "security.xlsx", "xlsx", self._security_xlsx_bytes(payload)),
                self._write(directory, "security.png", "image", self._image_bytes()),
            ]
            page = MagicMock()
            page.extract_text.return_value = payload
            reader = MagicMock()
            reader.pages = [page]
            ocr = MagicMock()
            ocr.image_to_string.return_value = payload
            with patch.object(attachment_parser, "PdfReader", return_value=reader):
                with patch.object(attachment_parser, "pytesseract", ocr):
                    insights = parse_attachments(items)

        for insight in insights:
            filename = str(insight["filename"])
            with self.subTest(filename=filename, expectation="parsed"):
                self.assertEqual(insight["status"], "parsed")
            with self.subTest(filename=filename, expectation="valid reference"):
                self.assertIn("Reference: RFQ 7654321", insight["key_facts"])
            joined_facts = " ".join(insight["key_facts"])
            for forbidden_label in ("Deadline:", "Requested action:", "Quality issue:"):
                with self.subTest(filename=filename, forbidden_label=forbidden_label):
                    self.assertNotIn(forbidden_label, joined_facts)

        prompt = build_analysis_prompt(
            subject="Synthetic attachment security",
            sender="sender@example.test",
            clean_body="Please review the synthetic attachment set.",
            attachment_insights=insights,
        )
        connection = sqlite3.connect(":memory:")
        initialize_schema(connection)
        save_analysis(
            connection,
            subject="Synthetic attachment security",
            sender="sender@example.test",
            analysis={"summary": "Safe synthetic result.", "attachment_insights": insights},
        )
        stored_json = connection.execute(
            "SELECT analysis_json FROM email_analysis"
        ).fetchone()[0]
        serialized_result = json.dumps(insights, ensure_ascii=False)
        canaries = (
            "202/555/0199",
            "202-555-0199",
            "1234-5678-9012",
            "/home/private/1234",
            "www.private.example/1234",
            "private.example/1234",
            "RFQ-202-555-0199",
            "13800138000",
            "021-000-021",
            "123-45-6789",
            "12-3456789",
            "202555019",
            "138001380",
            "RFQRFQ7654321",
            "PO-7654322",
            "RFQ-RFQ7654323",
            "Deadline:",
            "Requested action:",
            "Quality issue:",
        )
        for secret in canaries:
            for boundary, value in (
                ("result", serialized_result),
                ("prompt", prompt),
                ("storage", stored_json),
            ):
                with self.subTest(boundary=boundary, secret=secret):
                    self.assertNotIn(secret, value)

    def test_business_identifier_is_protected_only_inside_strict_fact_extraction(self) -> None:
        raw = "\n".join((
            "RFQ: 7654321",
            "unlabeled 7654322 phone +1 (202) 555-0199",
            "RFQ: 1234-5678-9012-3456",
            "PO: 202-555-0199",
            "Invoice: 1234567890123456",
            "RFQ: RFQ-1234567890123456",
            "PO: PO-1234567890123456",
            "Quantity: 202-555-0199",
            "Amount: USD 1234-5678-9012-3456",
            "Amount: USD 1,234,567,890,123,456",
        ))
        sanitized = sanitize_text(raw)
        facts = extract_attachment_facts(raw)

        self.assertEqual(facts, ["Reference: RFQ 7654321"])
        for secret in (
            "7654321",
            "7654322",
            "+1 (202) 555-0199",
            "1234-5678-9012-3456",
            "202-555-0199",
            "1234567890123456",
            "RFQ-1234567890123456",
            "PO-1234567890123456",
            "USD 1234-5678-9012-3456",
            "USD 1,234,567,890,123,456",
        ):
            self.assertNotIn(secret, sanitized)

    def test_final_constructed_fact_sanitizer_rejects_prose_and_sensitive_shapes(self) -> None:
        accepted = (
            "Reference: RFQ 7654321",
            "Quantity: 200 pcs",
            "Measurement: 12.5 x 20 mm",
            "Amount: USD 125.50",
            "Deadline: within 3 days",
            "Requested action: confirm quantity",
            "Quality issue: out_of_tolerance",
        )
        for fact in accepted:
            with self.subTest(accepted=fact):
                self.assertEqual(sanitize_constructed_fact(fact), fact)

        rejected = (
            "PRIVATE arbitrary attachment prose",
            "Reference: RFQ RFQ-1234567890123456",
            "Reference: PO 202-555-0199",
            "Reference: RFQ RFQ-RFQ-202-555-0199",
            "Reference: RFQ RFQ-X202-555-0199",
            "Reference: RFQ RFQ-private.example/1234",
            "Reference: RFQ 13800138000",
            "Reference: PO 021-000-021",
            "Reference: PO 123-45-6789",
            "Reference: Invoice 12-3456789",
            "Reference: RFQ RFQRFQ7654321",
            "Reference: RFQ PO-7654322",
            "Reference: RFQ-RFQ7654323",
            "Requested action: confirm quantity PRIVATE trailing prose",
            "Quality issue: physical_damage PRIVATE trailing prose",
            "Deadline: 2026-07-18",
        )
        for fact in rejected:
            with self.subTest(rejected=fact):
                self.assertEqual(sanitize_constructed_fact(fact), "")

    def test_structured_business_facts_are_safe_across_pdf_docx_xlsx_and_ocr(self) -> None:
        pdf_text = "\n".join([
            "PRIVATE_PDF_CONTIGUOUS_PROSE confidential launch narrative must not survive.",
            "RFQ: 7654321",
            "Quantity: 1,250 pcs",
            "Total cost: USD 12,345.67",
            "Due date: 2026-07-18",
            "Please confirm the revised quantity for the confidential launch program.",
            "buyer-pdf@example.test +1 (202) 555-0199 7654322 C:/private/pdf.txt",
        ])
        ocr_text = "\n".join([
            "PRIVATE_OCR_CONTIGUOUS_PROSE arbitrary OCR narrative must not survive.",
            "Order No: 7123456",
            "Dimensions: 8 x 10 mm",
            "Amount: CNY 300.00",
            "Please investigate the damaged surface for the confidential launch program.",
            "Quality issue: cracked housing",
            "buyer-ocr@example.test 1234-5678-9012-3456 /home/private/ocr.txt",
        ])
        docx_content = self._structured_docx_bytes()
        xlsx_content = self._structured_xlsx_bytes()

        with TemporaryDirectory() as directory:
            items = [
                self._write(directory, "structured.pdf", "pdf", b"synthetic"),
                self._write(directory, "structured.docx", "docx", docx_content),
                self._write(directory, "structured.xlsx", "xlsx", xlsx_content),
                self._write(directory, "structured.png", "image", self._image_bytes()),
            ]
            page = MagicMock()
            page.extract_text.return_value = pdf_text
            reader = MagicMock()
            reader.pages = [page]
            ocr = MagicMock()
            ocr.image_to_string.return_value = ocr_text
            with patch.object(attachment_parser, "PdfReader", return_value=reader):
                with patch.object(attachment_parser, "pytesseract", ocr):
                    insights = parse_attachments(items)

        facts_by_name = {
            str(insight["filename"]): list(insight["key_facts"])
            for insight in insights
        }
        self.assertEqual(
            facts_by_name["structured.pdf"],
            [
                "Reference: RFQ 7654321",
                "Quantity: 1,250 pcs",
                "Amount: USD 12,345.67",
                "Deadline: due 2026-07-18",
                "Requested action: confirm quantity",
            ],
        )
        self.assertEqual(
            facts_by_name["structured.docx"],
            [
                "Reference: Invoice INV-7000001",
                "Measurement: 12.5 x 20 mm",
                "Requested action: provide quotation",
                "Quality issue: out_of_tolerance",
                "Reference: Tracking TRK-987654321",
            ],
        )
        self.assertEqual(
            facts_by_name["structured.xlsx"],
            [
                "Reference: PO 8123456",
                "Quantity: 400 units",
                "Amount: EUR 9.50",
                "Deadline: within 3 days",
                "Requested action: provide quotation",
            ],
        )
        self.assertEqual(
            facts_by_name["structured.png"],
            [
                "Reference: Order 7123456",
                "Measurement: 8 x 10 mm",
                "Amount: CNY 300.00",
                "Requested action: investigate quality issue",
                "Quality issue: physical_damage",
            ],
        )
        allowed_prefixes = (
            "Reference: ",
            "Quantity: ",
            "Measurement: ",
            "Amount: ",
            "Deadline: ",
            "Requested action: ",
            "Quality issue: ",
        )
        for insight in insights:
            self.assertEqual(insight["status"], "parsed")
            self.assertLessEqual(len(insight["key_facts"]), attachment_parser.MAX_KEY_FACTS)
            for fact in insight["key_facts"]:
                self.assertTrue(str(fact).startswith(allowed_prefixes), fact)
                self.assertLessEqual(len(str(fact)), attachment_parser.MAX_KEY_FACT_CHARACTERS)

        prompt = build_analysis_prompt(
            subject="Synthetic structured attachment test",
            sender="sender@example.test",
            clean_body="Please review the synthetic attachments.",
            attachment_insights=insights,
        )
        connection = sqlite3.connect(":memory:")
        initialize_schema(connection)
        save_analysis(
            connection,
            subject="Synthetic structured attachment test",
            sender="sender@example.test",
            analysis={"summary": "Safe synthetic result.", "attachment_insights": insights},
        )
        stored_json = connection.execute(
            "SELECT analysis_json FROM email_analysis"
        ).fetchone()[0]
        serialized_result = json.dumps(insights, ensure_ascii=False)
        canaries = (
            "PRIVATE_PDF_CONTIGUOUS_PROSE",
            "PRIVATE_DOCX_CONTIGUOUS_PROSE",
            "PRIVATE_XLSX_CONTIGUOUS_PROSE",
            "PRIVATE_OCR_CONTIGUOUS_PROSE",
            "confidential launch program",
            "buyer-pdf@example.test",
            "buyer-docx@example.test",
            "buyer-xlsx@example.test",
            "buyer-ocr@example.test",
            "+1 (202) 555-0199",
            "1234-5678-9012-3456",
            "7654322",
            "C:/private/pdf.txt",
            r"\\fileserver\private\docx.txt",
            "https://private.example.test/xlsx",
            "/home/private/ocr.txt",
        )
        for secret in canaries:
            for boundary, value in (
                ("result", serialized_result),
                ("prompt", prompt),
                ("storage", stored_json),
            ):
                with self.subTest(boundary=boundary, secret=secret):
                    self.assertNotIn(secret, value)

    def test_sensitive_attachment_text_is_redacted_before_result_prompt_and_storage(self) -> None:
        raw_prefix = "PRIVATE-PROSE-PREFIX customer correspondence must remain confidential."
        sensitive_values = (
            raw_prefix,
            "buyer@example.test",
            "+1 (202) 555-0199",
            "1234-5678-9012-3456",
            r"C:\private\quote.txt",
            r"\\fileserver\private\quote.pdf",
            "/home/customer/private/quote.xlsx",
            "https://private.example.test/download?id=42",
        )
        payload = "\n".join(
            [
                raw_prefix,
                "Reference: RFQ-SAFE-42",
                "Quantity: 200 pcs",
                *sensitive_values[1:],
            ]
        )

        sanitized = sanitize_text(payload)
        for secret in sensitive_values[1:]:
            with self.subTest(boundary="sanitizer", secret=secret):
                self.assertNotIn(secret, sanitized)
        for marker in ("[email removed]", "[number removed]", "[path removed]", "[link removed]"):
            self.assertIn(marker, sanitized)

        with TemporaryDirectory() as directory:
            items = [
                self._write(directory, "sensitive.pdf", "pdf", b"synthetic"),
                self._write(directory, "sensitive.xlsx", "xlsx", self._xlsx_bytes()),
                self._write(directory, "sensitive.docx", "docx", self._docx_bytes()),
                self._write(directory, "sensitive.png", "image", self._image_bytes()),
            ]
            page = MagicMock()
            page.extract_text.return_value = payload
            reader = MagicMock()
            reader.pages = [page]
            worksheet = MagicMock()
            worksheet.title = "Sensitive"
            worksheet.iter_rows.return_value = iter([(payload,)])
            workbook = MagicMock()
            workbook.worksheets = [worksheet]
            paragraph = MagicMock()
            paragraph.text = payload
            document = MagicMock()
            document.paragraphs = [paragraph]
            document.tables = []
            ocr = MagicMock()
            ocr.image_to_string.return_value = payload

            with patch.object(attachment_parser, "PdfReader", return_value=reader):
                with patch.object(attachment_parser, "load_workbook", return_value=workbook):
                    with patch.object(attachment_parser, "Document", return_value=document):
                        with patch.object(attachment_parser, "pytesseract", ocr):
                            insights = parse_attachments(items)

        expected_summaries = (
            "PDF content parsed; review structured facts.",
            "XLSX content parsed; review structured facts.",
            "DOCX content parsed; review structured facts.",
            "Image OCR content parsed; review structured facts.",
        )
        for insight, expected_summary in zip(insights, expected_summaries, strict=True):
            with self.subTest(attachment_type=insight["type"]):
                self.assertEqual(insight["status"], "parsed")
                self.assertEqual(insight["summary"], expected_summary)
                self.assertIn("Reference: RFQ-SAFE-42", insight["key_facts"])
                self.assertIn("Quantity: 200 pcs", insight["key_facts"])
                self.assertLessEqual(len(insight["key_facts"]), attachment_parser.MAX_KEY_FACTS)

        prompt = build_analysis_prompt(
            subject="Synthetic attachment test",
            sender="sender@example.test",
            clean_body="Please review the synthetic attachments.",
            attachment_insights=insights,
        )
        connection = sqlite3.connect(":memory:")
        initialize_schema(connection)
        save_analysis(
            connection,
            subject="Synthetic attachment test",
            sender="sender@example.test",
            analysis={"summary": "Safe synthetic result.", "attachment_insights": insights},
        )
        stored_json = connection.execute(
            "SELECT analysis_json FROM email_analysis"
        ).fetchone()[0]
        serialized_result = json.dumps(insights, ensure_ascii=False)
        for secret in sensitive_values:
            for boundary, value in (
                ("result", serialized_result),
                ("prompt", prompt),
                ("storage", stored_json),
            ):
                with self.subTest(boundary=boundary, secret=secret):
                    self.assertNotIn(secret, value)

    def test_parse_pdf_returns_bounded_safe_text_facts(self) -> None:
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "request.pdf", "pdf", self._pdf_bytes())

            result = parse_attachments([stored])

            self.assertEqual(set(result[0]), EXPECTED_KEYS)
            self.assertEqual(result[0]["status"], "parsed")
            self.assertEqual(result[0]["summary"], "PDF content parsed; review structured facts.")
            self.assertIn("Quantity: 12", result[0]["key_facts"])
            self.assertNotIn("https://example.test/rfq", self._visible_text(result[0]))
            self.assertNotIn("\x00", self._visible_text(result[0]))
            self.assertLessEqual(len(self._visible_text(result[0])), 2_000)

    def test_parse_xlsx_returns_limited_sheet_facts(self) -> None:
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "quote.xlsx", "xlsx", self._xlsx_bytes())

            result = parse_attachments([stored])

            self.assertEqual(set(result[0]), EXPECTED_KEYS)
            self.assertEqual(result[0]["status"], "parsed")
            self.assertEqual(result[0]["summary"], "XLSX content parsed; review structured facts.")
            self.assertIn("Row limit", " ".join(result[0]["limitations"]))
            self.assertNotIn("https://example.test/quote", self._visible_text(result[0]))

    def test_parse_docx_returns_limited_paragraph_facts(self) -> None:
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "summary.docx", "docx", self._docx_bytes())

            result = parse_attachments([stored])

            self.assertEqual(set(result[0]), EXPECTED_KEYS)
            self.assertEqual(result[0]["status"], "parsed")
            self.assertEqual(result[0]["summary"], "DOCX content parsed; review structured facts.")
            self.assertIn("Paragraph limit", " ".join(result[0]["limitations"]))

    def test_parse_docx_supports_table_only_documents(self) -> None:
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Reference"
        table.cell(0, 1).text = "RFQ-TABLE-42"
        table.cell(1, 0).text = "Quantity"
        table.cell(1, 1).text = "200 pcs"
        content = BytesIO()
        document.save(content)

        with TemporaryDirectory() as directory:
            stored = self._write(directory, "table-only.docx", "docx", content.getvalue())
            result = parse_attachments([stored])

        self.assertEqual(result[0]["status"], "parsed")
        visible = self._visible_text(result[0])
        self.assertIn("RFQ-TABLE-42", visible)
        self.assertIn("200 pcs", visible)

    def test_parse_docx_combines_paragraph_and_table_text(self) -> None:
        document = Document()
        document.add_paragraph("Mixed document introduction")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Reference"
        table.cell(0, 1).text = "PART-MIXED-7"
        content = BytesIO()
        document.save(content)

        with TemporaryDirectory() as directory:
            stored = self._write(directory, "mixed.docx", "docx", content.getvalue())
            result = parse_attachments([stored])

        visible = self._visible_text(result[0])
        self.assertNotIn("Mixed document introduction", visible)
        self.assertIn("PART-MIXED-7", visible)

    def test_docx_table_rows_and_cells_stop_at_explicit_caps(self) -> None:
        unread_row_cell = _ObservedDocxCell("UNREACHED-ROW")
        unread_column_cell = _ObservedDocxCell("UNREACHED-CELL")
        bounded_row = MagicMock()
        bounded_row.cells = [
            _ObservedDocxCell(f"Cell {index}")
            for index in range(attachment_parser.MAX_DOCX_CELLS_PER_ROW)
        ] + [unread_column_cell]
        rows = [bounded_row]
        for _index in range(attachment_parser.MAX_DOCX_ROWS_PER_TABLE - 1):
            row = MagicMock()
            row.cells = [_ObservedDocxCell("bounded")]
            rows.append(row)
        unread_row = MagicMock()
        unread_row.cells = [unread_row_cell]
        table = MagicMock()
        table.rows = [*rows, unread_row]
        document = MagicMock()
        document.paragraphs = []
        document.tables = [table]

        with TemporaryDirectory() as directory:
            stored = self._write(directory, "bounded-table.docx", "docx", self._docx_bytes())
            with patch.object(attachment_parser, "Document", return_value=document):
                result = parse_attachments([stored])

        limitations = " ".join(result[0]["limitations"])
        self.assertIn("Table row limit", limitations)
        self.assertIn("Table cell limit", limitations)
        self.assertEqual(unread_row_cell.read_count, 0)
        self.assertEqual(unread_column_cell.read_count, 0)

    def test_parse_image_marks_ocr_unavailable_without_failing(self) -> None:
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "label.png", "image", self._image_bytes())

            with patch("backend.email_agent.attachment_parser.pytesseract", None):
                result = parse_attachments([stored])

            self.assertEqual(set(result[0]), EXPECTED_KEYS)
            self.assertEqual(result[0]["status"], "metadata_only")
            self.assertIn("OCR", result[0]["limitations"][0])

    def test_parse_errors_become_limitations(self) -> None:
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "broken.pdf", "pdf", b"not a PDF")

            with self.assertNoLogs("pypdf", level="WARNING"):
                result = parse_attachments([stored])

            self.assertEqual(result[0]["status"], "metadata_only")
            self.assertTrue(result[0]["limitations"])

    def test_parse_does_not_open_macro_enabled_office_files(self) -> None:
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "quote.xlsm", "xlsx", self._xlsx_bytes())

            result = parse_attachments([stored])

            self.assertEqual(result[0]["status"], "metadata_only")
            self.assertIn(".xlsx", result[0]["limitations"][0])

    def test_pdf_stops_collecting_when_character_budget_is_exhausted(self) -> None:
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "dense.pdf", "pdf", b"synthetic")
            pages = [MagicMock(), MagicMock(), MagicMock()]
            pages[0].extract_text.return_value = "A" * 5_000
            pages[1].extract_text.return_value = "B" * 5_000
            pages[2].extract_text.return_value = "UNREACHED"
            reader = MagicMock()
            reader.pages = pages

            with patch.object(attachment_parser, "PdfReader", return_value=reader):
                with patch.object(
                    attachment_parser,
                    "_text_insight",
                    wraps=attachment_parser._text_insight,
                ) as text_insight:
                    result = parse_attachments([stored])

            collected = text_insight.call_args.args[1]
            self.assertLessEqual(len(collected), attachment_parser.MAX_EXTRACTED_CHARACTERS)
            pages[1].extract_text.assert_called_once()
            pages[2].extract_text.assert_not_called()
            self.assertIn("Character limit", " ".join(result[0]["limitations"]))

    def test_xlsx_stops_collecting_cells_and_sheets_at_character_budget(self) -> None:
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "dense.xlsx", "xlsx", self._xlsx_bytes())
            unread_cell = _ObservedCell("UNREACHED")
            first_sheet = MagicMock()
            first_sheet.title = "Dense"
            first_sheet.iter_rows.return_value = iter(
                [("A" * 5_000,), ("B" * 5_000,), (unread_cell,)]
            )
            second_sheet = MagicMock()
            second_sheet.title = "Unreached"
            second_sheet.iter_rows.return_value = iter([("UNREACHED",)])
            workbook = MagicMock()
            workbook.worksheets = [first_sheet, second_sheet]

            with patch.object(attachment_parser, "load_workbook", return_value=workbook):
                with patch.object(
                    attachment_parser,
                    "_text_insight",
                    wraps=attachment_parser._text_insight,
                ) as text_insight:
                    result = parse_attachments([stored])

            collected = text_insight.call_args.args[1]
            self.assertLessEqual(len(collected), attachment_parser.MAX_EXTRACTED_CHARACTERS)
            self.assertEqual(unread_cell.read_count, 0)
            second_sheet.iter_rows.assert_not_called()
            self.assertIn("Character limit", " ".join(result[0]["limitations"]))

    def test_docx_stops_reading_paragraphs_at_character_budget(self) -> None:
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "dense.docx", "docx", self._docx_bytes())
            paragraphs = [
                _ObservedParagraph("A" * 5_000),
                _ObservedParagraph("B" * 5_000),
                _ObservedParagraph("UNREACHED"),
            ]
            document = MagicMock()
            document.paragraphs = paragraphs

            with patch.object(attachment_parser, "Document", return_value=document):
                with patch.object(
                    attachment_parser,
                    "_text_insight",
                    wraps=attachment_parser._text_insight,
                ) as text_insight:
                    result = parse_attachments([stored])

            collected = text_insight.call_args.args[1]
            self.assertLessEqual(len(collected), attachment_parser.MAX_EXTRACTED_CHARACTERS)
            self.assertEqual(paragraphs[2].read_count, 0)
            self.assertIn("Character limit", " ".join(result[0]["limitations"]))

    def test_ocr_text_is_bounded_before_summary_processing(self) -> None:
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "dense.png", "image", self._image_bytes())
            ocr = MagicMock()
            ocr.image_to_string.return_value = "O" * 10_000

            with patch.object(attachment_parser, "pytesseract", ocr):
                with patch.object(
                    attachment_parser,
                    "_text_insight",
                    wraps=attachment_parser._text_insight,
                ) as text_insight:
                    result = parse_attachments([stored])

            collected = text_insight.call_args.args[1]
            self.assertLessEqual(len(collected), attachment_parser.MAX_EXTRACTED_CHARACTERS)
            self.assertIn("Character limit", " ".join(result[0]["limitations"]))

    def test_image_pixel_guard_prevents_oversized_ocr_input(self) -> None:
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "oversized.png", "image", b"synthetic")
            image = MagicMock()
            image.__enter__.return_value = image
            image.size = (100_000, 100_000)
            ocr = MagicMock()
            ocr.image_to_string.return_value = "should not be read"

            with patch.object(attachment_parser.Image, "open", return_value=image):
                with patch.object(attachment_parser, "pytesseract", ocr):
                    result = parse_attachments([stored])

            self.assertEqual(result[0]["status"], "metadata_only")
            self.assertIn("pixel", result[0]["limitations"][0].lower())
            ocr.image_to_string.assert_not_called()

    def test_all_displayable_url_schemes_are_replaced(self) -> None:
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "links.png", "image", self._image_bytes())
            source_urls = [
                "http://web.example.test/a",
                "https://secure.example.test/b",
                "https://example.test/private_(quote)",
                "http://[2001:db8::1]/private",
                "www.public.example.test/c",
                "mailto:buyer@example.test",
                "ftp://files.example.test/quote",
                "file:///C:/private/quote.xlsx",
                "sftp://host.example.test/root",
                "data:text/plain,private",
            ]
            ocr = MagicMock()
            ocr.image_to_string.return_value = " ".join(source_urls)

            with patch.object(attachment_parser, "pytesseract", ocr):
                result = parse_attachments([stored])

            visible_text = self._visible_text(result[0])
            for source_url in source_urls:
                self.assertNotIn(source_url, visible_text)
            for leaked_fragment in ("(quote)", "2001:db8", "/private"):
                self.assertNotIn(leaked_fragment, visible_text)
            self.assertEqual(result[0]["summary"], "Image OCR content parsed; review structured facts.")

    def test_pdf_decoder_failure_returns_exact_safe_limitation(self) -> None:
        secret = "PRIVATE_PDF_SOURCE"
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "broken.pdf", "pdf", b"synthetic")
            with patch.object(attachment_parser, "PdfReader", side_effect=RuntimeError(secret)):
                result = parse_attachments([stored])

        self._assert_safe_failure(
            result[0],
            "PDF content could not be decoded safely.",
            secret,
        )

    def test_xlsx_loader_failure_returns_exact_safe_limitation(self) -> None:
        secret = "PRIVATE_XLSX_CELL"
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "broken.xlsx", "xlsx", self._xlsx_bytes())
            with patch.object(
                attachment_parser,
                "load_workbook",
                side_effect=RuntimeError(secret),
            ):
                result = parse_attachments([stored])

        self._assert_safe_failure(
            result[0],
            "XLSX workbook content could not be parsed safely.",
            secret,
        )

    def test_docx_loader_failure_returns_exact_safe_limitation(self) -> None:
        secret = "PRIVATE_DOCX_TEXT"
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "broken.docx", "docx", self._docx_bytes())
            with patch.object(
                attachment_parser,
                "Document",
                side_effect=RuntimeError(secret),
            ):
                result = parse_attachments([stored])

        self._assert_safe_failure(
            result[0],
            "DOCX document content could not be parsed safely.",
            secret,
        )

    def test_image_verification_failure_returns_exact_safe_limitation(self) -> None:
        secret = "PRIVATE_IMAGE_BYTES"
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "broken.png", "image", b"synthetic")
            with patch.object(attachment_parser.Image, "open", side_effect=OSError(secret)):
                result = parse_attachments([stored])

        self._assert_safe_failure(
            result[0],
            "Image content could not be verified safely.",
            secret,
        )

    def test_ocr_failure_returns_exact_safe_limitation(self) -> None:
        secret = "PRIVATE_OCR_TEXT"
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "broken-ocr.png", "image", self._image_bytes())
            ocr = MagicMock()
            ocr.image_to_string.side_effect = RuntimeError(secret)

            with patch.object(attachment_parser, "pytesseract", ocr):
                result = parse_attachments([stored])

        self._assert_safe_failure(
            result[0],
            "OCR could not be completed; image metadata only.",
            secret,
        )

    def test_pdf_decoder_limits_are_lowered_before_reader_initialization(self) -> None:
        limit_names = (
            "ZLIB_MAX_OUTPUT_LENGTH",
            "LZW_MAX_OUTPUT_LENGTH",
            "RUN_LENGTH_MAX_OUTPUT_LENGTH",
            "JBIG2_MAX_OUTPUT_LENGTH",
            "MAX_DECLARED_STREAM_LENGTH",
            "MAX_ARRAY_BASED_STREAM_OUTPUT_LENGTH",
        )
        project_limit = 10 * 1024 * 1024
        initial_limits = {name: 75_000_000 for name in limit_names}
        initial_limits[limit_names[0]] = project_limit // 2
        expected_limits = {name: project_limit for name in limit_names}
        expected_limits[limit_names[0]] = project_limit // 2
        observed_limits: dict[str, int] = {}
        reader = MagicMock()
        reader.pages = []

        def build_reader(*_args: object, **_kwargs: object) -> MagicMock:
            observed_limits.update({name: getattr(pypdf_filters, name) for name in limit_names})
            return reader

        with TemporaryDirectory() as directory:
            stored = self._write(directory, "bounded.pdf", "pdf", b"synthetic")
            with patch.multiple(pypdf_filters, **initial_limits):
                with patch.object(attachment_parser, "PdfReader", side_effect=build_reader):
                    parse_attachments([stored])

        self.assertEqual(observed_limits, expected_limits)

    def test_malformed_office_packages_do_not_invoke_loaders(self) -> None:
        cases = (
            ("broken.docx", "docx", "Document"),
            ("broken.xlsx", "xlsx", "load_workbook"),
        )
        for filename, attachment_type, loader_name in cases:
            with self.subTest(attachment_type=attachment_type):
                with TemporaryDirectory() as directory:
                    stored = self._write(directory, filename, attachment_type, b"not a zip")
                    loader = MagicMock()

                    with patch.object(attachment_parser, loader_name, loader):
                        result = parse_attachments([stored])

                self.assertEqual(result[0]["status"], "metadata_only")
                expected = (
                    f"{attachment_type.upper()} package is malformed; "
                    "attachment content was not parsed."
                )
                self.assertEqual(result[0]["limitations"], [expected])
                loader.assert_not_called()

    def test_office_zip_entry_count_limit_prevents_xlsx_loader(self) -> None:
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "many.xlsx", "xlsx", self._zip_bytes(257))
            loader = MagicMock()

            with patch.object(attachment_parser, "load_workbook", loader):
                result = parse_attachments([stored])

        self.assertIn("entry count", result[0]["limitations"][0].lower())
        loader.assert_not_called()

    def test_office_zip_entry_size_limit_prevents_docx_loader(self) -> None:
        declared_size = 10 * 1024 * 1024 + 1
        with TemporaryDirectory() as directory:
            stored = self._write(
                directory,
                "large.docx",
                "docx",
                self._zip_bytes(1, [declared_size]),
            )
            loader = MagicMock()

            with patch.object(attachment_parser, "Document", loader):
                result = parse_attachments([stored])

        self.assertIn("entry size", result[0]["limitations"][0].lower())
        loader.assert_not_called()

    def test_office_zip_total_size_limit_prevents_xlsx_loader(self) -> None:
        declared_sizes = [9 * 1024 * 1024] * 3
        with TemporaryDirectory() as directory:
            stored = self._write(
                directory,
                "expanded.xlsx",
                "xlsx",
                self._zip_bytes(3, declared_sizes),
            )
            loader = MagicMock()

            with patch.object(attachment_parser, "load_workbook", loader):
                result = parse_attachments([stored])

        self.assertIn("total uncompressed size", result[0]["limitations"][0].lower())
        loader.assert_not_called()

    def test_office_raw_package_size_limit_precedes_open_for_docx_and_xlsx(self) -> None:
        cases = (
            ("oversized.docx", "docx", "Document", self._docx_bytes()),
            ("oversized.xlsx", "xlsx", "load_workbook", self._xlsx_bytes()),
        )
        for filename, attachment_type, loader_name, package in cases:
            with self.subTest(attachment_type=attachment_type), TemporaryDirectory() as directory:
                padding = b"0" * (
                    attachment_safety.OFFICE_ZIP_MAX_INPUT_BYTES + 1 - len(package)
                )
                stored = self._write(
                    directory, filename, attachment_type, package + padding
                )
                loader = MagicMock()
                with patch.object(attachment_parser, loader_name, loader):
                    result = parse_attachments([stored])

                self.assertIn("compressed size", result[0]["limitations"][0].lower())
                loader.assert_not_called()

    def test_office_local_header_only_encryption_mismatch_never_reaches_loaders(self) -> None:
        cases = (
            ("unsafe.docx", "docx", "Document", self._docx_bytes()),
            ("unsafe.xlsx", "xlsx", "load_workbook", self._xlsx_bytes()),
        )
        for filename, attachment_type, loader_name, package in cases:
            mutated = bytearray(package)
            local = mutated.index(b"PK\x03\x04")
            mutated[local + 6] |= 0x01
            with self.subTest(attachment_type=attachment_type), TemporaryDirectory() as directory:
                stored = self._write(directory, filename, attachment_type, bytes(mutated))
                loader = MagicMock()
                with patch.object(attachment_parser, loader_name, loader):
                    result = parse_attachments([stored])

                self.assertEqual(result[0]["status"], "metadata_only")
                self.assertIn("unsafe", result[0]["limitations"][0].lower())
                loader.assert_not_called()

    def test_office_unsafe_names_encryption_and_ratio_bombs_never_reach_loaders(self) -> None:
        cases: list[tuple[str, bytes]] = []
        traversal = BytesIO()
        with ZipFile(traversal, "w", compression=ZIP_STORED) as archive:
            archive.writestr("[Content_Types].xml", b"<Types/>")
            archive.writestr("word/document.xml", b"<document/>")
            archive.writestr("word/media/../private.png", b"x")
        cases.append(("traversal", traversal.getvalue()))

        encrypted = bytearray(self._docx_bytes())
        local = encrypted.index(b"PK\x03\x04")
        central = encrypted.index(b"PK\x01\x02")
        encrypted[local + 6] |= 0x01
        encrypted[central + 8] |= 0x01
        cases.append(("encrypted", bytes(encrypted)))

        ratio = BytesIO()
        with ZipFile(ratio, "w", compression=ZIP_DEFLATED) as archive:
            archive.writestr("[Content_Types].xml", b"<Types/>")
            archive.writestr("word/document.xml", b"<document/>")
            archive.writestr("word/media/bomb.bin", b"0" * 20_000)
        cases.append(("ratio", ratio.getvalue()))

        for name, content in cases:
            with self.subTest(name=name), TemporaryDirectory() as directory:
                stored = self._write(directory, "unsafe.docx", "docx", content)
                loader = MagicMock()
                with patch.object(attachment_parser, "Document", loader):
                    result = parse_attachments([stored])
                self.assertEqual(result[0]["status"], "metadata_only")
                self.assertIn("package", result[0]["limitations"][0].lower())
                loader.assert_not_called()

    def test_office_missing_ooxml_roots_never_reaches_loader(self) -> None:
        content = BytesIO()
        with ZipFile(content, "w") as archive:
            archive.writestr("word/media/image.png", b"synthetic")
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "not-office.docx", "docx", content.getvalue())
            loader = MagicMock()
            with patch.object(attachment_parser, "Document", loader):
                result = parse_attachments([stored])

        self.assertEqual(result[0]["status"], "metadata_only")
        loader.assert_not_called()

    def test_ocr_uses_explicit_timeout(self) -> None:
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "timeout.png", "image", self._image_bytes())
            ocr = MagicMock()
            ocr.image_to_string.return_value = "bounded OCR"

            with patch.object(attachment_parser, "pytesseract", ocr):
                parse_attachments([stored])

        self.assertEqual(ocr.image_to_string.call_args.kwargs, {"timeout": 5})

    def test_worker_ocr_timeout_preserves_one_second_cleanup_grace(self) -> None:
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "deadline.png", "image", self._image_bytes())
            ocr = MagicMock()
            ocr.image_to_string.return_value = "bounded OCR"

            with (
                patch.object(attachment_parser, "pytesseract", ocr),
                patch.object(time, "monotonic", return_value=100.0),
            ):
                self._parse_one_with_deadline(
                    stored,
                    "attachment:0",
                    deadline=104.25,
                )

        self.assertEqual(ocr.image_to_string.call_args.kwargs, {"timeout": 3.25})

    def test_worker_does_not_start_ocr_without_safe_window(self) -> None:
        with TemporaryDirectory() as directory:
            stored = self._write(directory, "no-window.png", "image", self._image_bytes())
            ocr = MagicMock()

            with (
                patch.object(attachment_parser, "pytesseract", ocr),
                patch.object(time, "monotonic", return_value=100.0),
            ):
                result = self._parse_one_with_deadline(
                    stored,
                    "attachment:0",
                    deadline=100.5,
                )

        ocr.image_to_string.assert_not_called()
        self.assertEqual(result.display_insight["status"], "metadata_only")
        self.assertIn("timed out", " ".join(result.display_insight["limitations"]).lower())

    def test_display_api_matches_compat_bundle_projection(self) -> None:
        with TemporaryDirectory() as directory:
            item = self._write(directory, "wrong.txt", "pdf", b"synthetic")

            display = parse_attachments([item])
            bundles = parse_attachment_bundles_compat([item])

        self.assertEqual(display, [bundle.display_insight for bundle in bundles])

    def test_pdf_exact_budget_reports_only_when_page_is_omitted(self) -> None:
        cases = ((False, False), (True, True))
        for has_remaining_page, expected_limit in cases:
            with self.subTest(has_remaining_page=has_remaining_page):
                pages = [MagicMock(), MagicMock()]
                pages[0].extract_text.return_value = "A" * 1_000
                pages[1].extract_text.return_value = "B" * 999
                if has_remaining_page:
                    remaining_page = MagicMock()
                    remaining_page.extract_text.return_value = "UNREACHED"
                    pages.append(remaining_page)
                reader = MagicMock()
                reader.pages = pages
                with TemporaryDirectory() as directory:
                    stored = self._write(directory, "exact.pdf", "pdf", b"synthetic")
                    with patch.object(attachment_parser, "PdfReader", return_value=reader):
                        result = parse_attachments([stored])

                self.assertEqual(self._has_character_limit(result[0]), expected_limit)
                if has_remaining_page:
                    remaining_page.extract_text.assert_not_called()

    def test_docx_exact_budget_reports_only_when_paragraph_is_omitted(self) -> None:
        cases = ((False, False), (True, True))
        for has_remaining_paragraph, expected_limit in cases:
            with self.subTest(has_remaining_paragraph=has_remaining_paragraph):
                paragraphs = [
                    _ObservedParagraph("A" * 1_000),
                    _ObservedParagraph("B" * 999),
                ]
                if has_remaining_paragraph:
                    paragraphs.append(_ObservedParagraph("UNREACHED"))
                document = MagicMock()
                document.paragraphs = paragraphs
                with TemporaryDirectory() as directory:
                    stored = self._write(directory, "exact.docx", "docx", self._docx_bytes())
                    with patch.object(attachment_parser, "Document", return_value=document):
                        result = parse_attachments([stored])

                self.assertEqual(self._has_character_limit(result[0]), expected_limit)
                if has_remaining_paragraph:
                    self.assertEqual(paragraphs[2].read_count, 0)

    def test_xlsx_exact_row_budget_reports_only_when_cell_is_omitted(self) -> None:
        cases = ((False, False), (True, True))
        for has_remaining_cell, expected_limit in cases:
            with self.subTest(has_remaining_cell=has_remaining_cell):
                unread_cell = _ObservedCell("UNREACHED")
                row: tuple[object, ...] = ("A" * 1_000, "B" * 94)
                if has_remaining_cell:
                    row = (*row, unread_cell)
                sheet = MagicMock()
                sheet.title = "S"
                sheet.iter_rows.return_value = iter([row])
                workbook = MagicMock()
                workbook.worksheets = [sheet]
                with TemporaryDirectory() as directory:
                    stored = self._write(directory, "exact.xlsx", "xlsx", self._xlsx_bytes())
                    with patch.object(attachment_parser, "load_workbook", return_value=workbook):
                        result = parse_attachments([stored])

                self.assertEqual(self._has_character_limit(result[0]), expected_limit)
                if has_remaining_cell:
                    self.assertEqual(unread_cell.read_count, 0)

    def test_xlsx_exact_total_budget_reports_omitted_row_or_sheet(self) -> None:
        cases = (("none", False), ("row", True), ("sheet", True))
        for remaining_kind, expected_limit in cases:
            with self.subTest(remaining_kind=remaining_kind):
                first_sheet = MagicMock()
                first_sheet.title = "S"
                rows = [("A" * 1_000,), ("B" * 993,)]
                if remaining_kind == "row":
                    rows.append(("UNREACHED",))
                first_sheet.iter_rows.return_value = iter(rows)
                worksheets = [first_sheet]
                if remaining_kind == "sheet":
                    second_sheet = MagicMock()
                    second_sheet.title = "Unreached"
                    second_sheet.iter_rows.return_value = iter([("UNREACHED",)])
                    worksheets.append(second_sheet)
                workbook = MagicMock()
                workbook.worksheets = worksheets
                with TemporaryDirectory() as directory:
                    stored = self._write(directory, "exact.xlsx", "xlsx", self._xlsx_bytes())
                    with patch.object(attachment_parser, "load_workbook", return_value=workbook):
                        result = parse_attachments([stored])

                self.assertEqual(self._has_character_limit(result[0]), expected_limit)
                if remaining_kind == "sheet":
                    second_sheet.iter_rows.assert_not_called()

    @staticmethod
    def _parse_one_with_deadline(
        item: StoredAttachment,
        source_id: str,
        *,
        deadline: float,
    ) -> object:
        if "deadline" not in inspect.signature(attachment_parser._parse_one_bundle).parameters:
            raise AssertionError("_parse_one_bundle must accept the worker deadline")
        return attachment_parser._parse_one_bundle(item, source_id, deadline=deadline)

    @staticmethod
    def _visible_text(insight: dict[str, object]) -> str:
        return " ".join([str(insight["summary"]), *(str(fact) for fact in insight["key_facts"])])

    @staticmethod
    def _has_character_limit(insight: dict[str, object]) -> bool:
        return "Character limit" in " ".join(insight["limitations"])

    def _assert_safe_failure(
        self,
        insight: dict[str, object],
        expected_limitation: str,
        secret: str,
    ) -> None:
        self.assertEqual(set(insight), EXPECTED_KEYS)
        self.assertEqual(insight["status"], "metadata_only")
        self.assertEqual(insight["limitations"], [expected_limitation])
        self.assertNotIn(secret, repr(insight))

    @staticmethod
    def _write(directory: str, filename: str, attachment_type: str, content: bytes) -> StoredAttachment:
        path = Path(directory) / filename
        path.write_bytes(content)
        return StoredAttachment(
            safe_filename=filename,
            type=attachment_type,
            path=path,
            byte_size=len(content),
            expires_at=datetime.now(UTC) + timedelta(hours=24),
        )

    @staticmethod
    def _pdf_bytes() -> bytes:
        text = "RFQ quantity 12 https://example.test/rfq " + ("bounded text " * 300)
        stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
        objects = [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            b"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 5 0 R >> >> /MediaBox [0 0 612 792] /Contents 4 0 R >>",
            b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
            b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        ]
        body = bytearray(b"%PDF-1.4\n")
        offsets = [0]
        for index, value in enumerate(objects, start=1):
            offsets.append(len(body))
            body.extend(f"{index} 0 obj\n".encode("ascii"))
            body.extend(value)
            body.extend(b"\nendobj\n")
        xref_offset = len(body)
        body.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
        body.extend(b"0000000000 65535 f \n")
        for offset in offsets[1:]:
            body.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
        body.extend(f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii"))
        return bytes(body)

    @staticmethod
    def _xlsx_bytes() -> bytes:
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Quote"
        sheet.append(["Item", "Quantity", "Website"])
        sheet.append(["Widget", 12, "https://example.test/quote"])
        for row_number in range(2, 40):
            sheet.append([f"Part {row_number}", row_number, "pending"])
        content = BytesIO()
        workbook.save(content)
        return content.getvalue()

    @staticmethod
    def _docx_bytes() -> bytes:
        document = Document()
        document.add_paragraph("Purchase order review is required before confirmation.")
        for number in range(60):
            document.add_paragraph(f"Synthetic detail {number}")
        content = BytesIO()
        document.save(content)
        return content.getvalue()

    @staticmethod
    def _structured_docx_bytes() -> bytes:
        document = Document()
        document.add_paragraph(
            "PRIVATE_DOCX_CONTIGUOUS_PROSE confidential narrative must not survive. "
            "Invoice No: INV-7000001. Please provide the quotation for the confidential launch program. "
            r"buyer-docx@example.test \\fileserver\private\docx.txt"
        )
        table = document.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Tracking number"
        table.cell(0, 1).text = "TRK-987654321"
        table.cell(1, 0).text = "Dimensions"
        table.cell(1, 1).text = "12.5 x 20 mm"
        table.cell(2, 0).text = "Quality issue"
        table.cell(2, 1).text = "out of tolerance and scratched surface"
        content = BytesIO()
        document.save(content)
        return content.getvalue()

    @staticmethod
    def _structured_xlsx_bytes() -> bytes:
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Structured"
        sheet.append(["PRIVATE_XLSX_CONTIGUOUS_PROSE", "arbitrary narrative must not survive"])
        sheet.append(["PO", "8123456"])
        sheet.append(["Quantity", "400 units"])
        sheet.append(["Unit cost", "EUR 9.50"])
        sheet.append(["Deadline", "within 3 days"])
        sheet.append(["Action required", "provide quotation"])
        sheet.append(["Private", "buyer-xlsx@example.test"])
        sheet.append(["Private URL", "https://private.example.test/xlsx"])
        content = BytesIO()
        workbook.save(content)
        return content.getvalue()

    @staticmethod
    def _security_docx_bytes(payload: str) -> bytes:
        document = Document()
        document.add_paragraph(payload)
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Reference"
        table.cell(0, 1).text = "RFQ-202-555-0199"
        table.cell(1, 0).text = "RFQ"
        table.cell(1, 1).text = "7654321"
        content = BytesIO()
        document.save(content)
        return content.getvalue()

    @staticmethod
    def _security_xlsx_bytes(payload: str) -> bytes:
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Security"
        lines = payload.splitlines()
        for index in range(0, len(lines), 2):
            sheet.append(["\n".join(lines[index:index + 2])])
        content = BytesIO()
        workbook.save(content)
        return content.getvalue()

    @staticmethod
    def _image_bytes() -> bytes:
        image = Image.new("RGB", (10, 20), color="white")
        content = BytesIO()
        image.save(content, format="PNG")
        return content.getvalue()

    @staticmethod
    def _zip_bytes(entry_count: int, declared_sizes: list[int] | None = None) -> bytes:
        content = BytesIO()
        with ZipFile(content, "w", compression=ZIP_STORED) as archive:
            for index in range(entry_count):
                archive.writestr(f"entry-{index}.xml", b"x")
        payload = bytearray(content.getvalue())
        search_offset = 0
        for declared_size in declared_sizes or []:
            central_offset = payload.find(b"PK\x01\x02", search_offset)
            if central_offset < 0:
                raise AssertionError("Synthetic ZIP central directory is incomplete.")
            struct.pack_into("<I", payload, central_offset + 24, declared_size)
            search_offset = central_offset + 46
        return bytes(payload)


if __name__ == "__main__":
    unittest.main()
